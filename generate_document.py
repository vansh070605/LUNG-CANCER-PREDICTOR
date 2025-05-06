from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def add_screenshot_placeholder(doc, title, description):
    doc.add_heading(title, 3)
    doc.add_paragraph(description)
    doc.add_paragraph("Screenshot Placeholder: [Insert screenshot here]", style='Intense Quote')
    doc.add_paragraph()

def create_title_page(doc):
    # Add title
    title = doc.add_heading('LUNG CANCER PREDICTOR', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add subtitle
    subtitle = doc.add_paragraph('A Modern Web Application for Lung Cancer Risk Assessment')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date
    date = doc.add_paragraph(f'Generated on: {datetime.datetime.now().strftime("%B %d, %Y")}')
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()

def create_executive_summary(doc):
    doc.add_heading('Executive Summary', 1)
    summary = doc.add_paragraph()
    summary.add_run('The Lung Cancer Predictor is a robust software solution designed to streamline and optimize the process of lung cancer risk assessment. This project aims to provide healthcare professionals and patients with an efficient and intuitive platform for managing medical histories, risk assessments, predictions, and patient interactions. Built with Flask and MySQL, it features a modern glassmorphism UI design and implements robust security measures. The application provides real-time risk assessment, maintains comprehensive medical histories, and ensures data integrity through advanced concurrency control mechanisms.')
    doc.add_page_break()

def create_project_overview(doc):
    doc.add_heading('Project Overview', 1)
    
    # Purpose
    doc.add_heading('Purpose', 2)
    purpose = doc.add_paragraph('The primary goal of this project is to provide an accessible and user-friendly platform for lung cancer risk assessment. The application helps users understand their risk factors and provides valuable medical recommendations based on their input. The DBMS serves as the foundational framework for storing, organizing, and processing data related to patient information, medical histories, risk assessments, and prediction results. By centralizing data management and providing efficient data retrieval mechanisms, the system facilitates more accurate risk assessments, improves patient care, and enables data-driven medical decision-making.')
    
    # Key Objectives
    doc.add_heading('Key Objectives', 2)
    objectives = [
        'Provide accurate lung cancer risk assessment based on user symptoms and medical history',
        'Ensure secure user authentication and data protection',
        'Implement real-time risk score calculation',
        'Maintain comprehensive medical history tracking',
        'Offer user-friendly interface with modern design principles',
        'Enable efficient data management and retrieval',
        'Support concurrent user access with data integrity'
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')
    
    # Advantages
    doc.add_heading('Key Advantages', 2)
    advantages = [
        'Data Accuracy: The system ensures data accuracy and consistency by enforcing data integrity constraints, minimizing errors in risk assessment and medical history tracking.',
        'Scalability: The DBMS supports scalability, allowing the system to handle increased user demand, additional risk factors, and expanded medical data without compromising performance.',
        'Security: The system implements secure multi-user access, concurrency control, and data privacy measures, ensuring reliable and efficient operations even during peak usage.',
        'Real-time Processing: Enables immediate risk assessment calculations and updates to medical histories.',
        'Comprehensive Tracking: Maintains detailed records of patient histories, risk assessments, and medical recommendations.'
    ]
    for advantage in advantages:
        doc.add_paragraph(advantage, style='List Bullet')
    
    doc.add_page_break()

def create_technical_architecture(doc):
    doc.add_heading('Technical Architecture', 1)
    
    # Backend
    doc.add_heading('Backend Technologies', 2)
    backend = [
        'Python 3.12',
        'Flask 3.0.2',
        'MySQL 8.0',
        'Flask-Login',
        'Flask-Bcrypt',
        'PyJWT'
    ]
    for tech in backend:
        doc.add_paragraph(tech, style='List Bullet')
    
    # Frontend
    doc.add_heading('Frontend Technologies', 2)
    frontend = [
        'HTML5',
        'CSS3',
        'JavaScript',
        'Glassmorphism Design'
    ]
    for tech in frontend:
        doc.add_paragraph(tech, style='List Bullet')
    
    doc.add_page_break()

def create_features_section(doc):
    doc.add_heading('Features and Functionality', 1)
    
    features = {
        'User Authentication': [
            'Secure registration and login',
            'Password hashing',
            'Session management',
            'JWT token support'
        ],
        'Risk Assessment': [
            'Age and gender-based analysis',
            'Smoking status evaluation',
            'Symptom assessment',
            'Real-time risk score calculation'
        ],
        'Medical History': [
            'Family history tracking',
            'Smoking history',
            'Previous lung diseases',
            'Occupational exposure'
        ],
        'Data Management': [
            'Prediction history',
            'User feedback system',
            'Medical recommendations',
            'Symptom database'
        ]
    }
    
    for feature, details in features.items():
        doc.add_heading(feature, 2)
        for detail in details:
            doc.add_paragraph(detail, style='List Bullet')
    
    doc.add_page_break()

def create_database_design(doc):
    doc.add_heading('Database Design', 1)
    
    # Database Schema
    doc.add_heading('Database Schema', 2)
    tables = {
        'users': 'User information and authentication details',
        'medical_history': 'Medical background and history',
        'predictions': 'Risk assessment results',
        'user_predictions': 'User-prediction mapping',
        'symptoms': 'Symptom database',
        'recommendations': 'Medical advice and recommendations',
        'user_feedback': 'User feedback and ratings'
    }
    
    for table, description in tables.items():
        p = doc.add_paragraph()
        p.add_run(f'{table}: ').bold = True
        p.add_run(description)
    
    # Key Features
    doc.add_heading('Key Database Features', 2)
    features = [
        'Lock Management: Only one user can make a prediction at a time',
        'Version Control: Prevents lost updates and supports optimistic concurrency',
        'Transaction Log: All changes are logged for audit and recovery',
        'Backup & Recovery: Daily backups and point-in-time recovery procedures',
        'Deadlock Detection: Automatic cleanup of expired locks'
    ]
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_page_break()

def create_ui_section(doc):
    doc.add_heading('User Interface', 1)
    
    features = [
        'Modern glassmorphism design',
        'Dark theme with red accents',
        'Responsive layout',
        'Interactive forms',
        'Real-time feedback',
        'Animated transitions',
        'Mobile-friendly design'
    ]
    
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')
    
    # Add screenshot placeholders
    doc.add_heading('UI Screenshots', 2)
    screenshots = [
        ('Login Page', 'The secure login interface with glassmorphism design'),
        ('Registration Form', 'User registration form with input validation'),
        ('Dashboard', 'Main dashboard showing user information and quick actions'),
        ('Prediction Form', 'Interactive form for entering symptoms and medical history'),
        ('Results Page', 'Detailed risk assessment results with recommendations'),
        ('Mobile View', 'Responsive design on mobile devices')
    ]
    
    for title, description in screenshots:
        add_screenshot_placeholder(doc, title, description)
    
    doc.add_page_break()

def create_security_section(doc):
    doc.add_heading('Security Features', 1)
    
    features = [
        'Password hashing using Flask-Bcrypt',
        'JWT token authentication',
        'SQL injection prevention',
        'Input validation',
        'Session management',
        'Protected routes'
    ]
    
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_page_break()

def create_implementation_details(doc):
    doc.add_heading('Implementation Details', 1)
    
    # Concurrency Control
    doc.add_heading('Concurrency Control', 2)
    doc.add_paragraph('The application implements a global locking mechanism to ensure that only one user can make a prediction at a time. This prevents race conditions and ensures data consistency.')
    
    # Error Handling
    doc.add_heading('Error Handling', 2)
    doc.add_paragraph('Comprehensive error handling is implemented throughout the application, with user-friendly error messages and proper logging of system errors.')
    
    # Add screenshot placeholders for error messages
    doc.add_heading('Error Handling Screenshots', 2)
    error_screenshots = [
        ('Concurrency Error', 'Error message displayed when another user is making a prediction'),
        ('Validation Error', 'Input validation error message example'),
        ('Authentication Error', 'Login/registration error message example')
    ]
    
    for title, description in error_screenshots:
        add_screenshot_placeholder(doc, title, description)
    
    doc.add_page_break()

def create_future_enhancements(doc):
    doc.add_heading('Future Enhancements', 1)
    
    enhancements = [
        'Integration with medical imaging data',
        'Machine learning model improvements',
        'Mobile application development',
        'Integration with electronic health records',
        'Multi-language support',
        'Advanced analytics dashboard'
    ]
    
    for enhancement in enhancements:
        doc.add_paragraph(enhancement, style='List Bullet')

def main():
    # Create document
    doc = Document()
    
    # Set document properties
    doc.core_properties.title = 'Lung Cancer Predictor Documentation'
    doc.core_properties.author = 'Project Team'
    
    # Create sections
    create_title_page(doc)
    create_executive_summary(doc)
    create_project_overview(doc)
    create_technical_architecture(doc)
    create_features_section(doc)
    create_database_design(doc)
    create_ui_section(doc)
    create_security_section(doc)
    create_implementation_details(doc)
    create_future_enhancements(doc)
    
    # Add instructions for screenshots
    doc.add_heading('Instructions for Adding Screenshots', 1)
    instructions = [
        'Take screenshots of each major feature and interface element',
        'Save screenshots in a high-resolution format (PNG or JPG)',
        'Replace the placeholder text with actual screenshots',
        'Ensure screenshots are clear and properly cropped',
        'Add captions to explain what each screenshot demonstrates'
    ]
    for instruction in instructions:
        doc.add_paragraph(instruction, style='List Bullet')
    
    # Save document
    doc.save('Lung_Cancer_Predictor_Documentation.docx')

if __name__ == '__main__':
    main() 